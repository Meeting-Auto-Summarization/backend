from docx import Document
from docx.shared import Inches
import sys, json

if __name__ == '__main__':
    meeting = json.loads(sys.argv[1])
    report = json.loads(sys.argv[2])

    document = Document()

    document.add_heading(meeting['title'], 0)

    for i in range(len(report)):
        for j in range(len(report[i])):
            if j == 0:
                document.add_heading(report[i][j]['title'], level=1)
            else:
                heading = document.add_heading(report[i][j]['title'], level=2)
                p = document.add_paragraph(report[i][j]['summary'], style='List Bullet')

    id = meeting['_id']
    document.save(f'{id}.docx')
    print(id)
    sys.stdout.flush()
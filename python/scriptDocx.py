from docx import Document
from docx.shared import Inches
import sys, json

if __name__ == '__main__':
    meeting = json.loads(sys.argv[1])
    script = json.loads(sys.argv[2])

    document = Document()

    document.add_heading(meeting['title'], 0)

    table = document.add_table(rows=len(script), cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Time'
    hdr_cells[2].text = 'Content'
    for i in range(len(script)):
        row_cells = table.add_row().cells
        row_cells[0].text = script[i]['nick']

        stime = int(script[i]['time'])

        hour = int(stime / 3600)
        min = int((stime / 60) % 60)
        sec = int(stime % 60)

        if hour == 0:
            stime = f'{min}:{sec}'
        else:
            stime = f'{hour}:{min}:{sec}'
        
        row_cells[1].text = stime
        row_cells[2].text = script[i]['content']

    id = meeting['_id']
    document.save(f'{id}.docx')
    print(id)
    sys.stdout.flush()